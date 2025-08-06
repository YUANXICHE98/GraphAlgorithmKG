"""
文档处理模块
支持多种格式文档的解析和预处理
"""

import json
import re
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from pathlib import Path
import mimetypes

# 文档解析库
try:
    import PyPDF2
    import docx
    import pandas as pd
    from bs4 import BeautifulSoup
    import markdown
except ImportError as e:
    print(f"警告: 某些文档解析库未安装: {e}")

@dataclass
class DocumentMetadata:
    """文档元数据"""
    filename: str
    file_type: str
    file_size: int
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    language: str = "zh"
    encoding: str = "utf-8"
    created_at: str = ""
    modified_at: str = ""

@dataclass
class ProcessedDocument:
    """处理后的文档"""
    content: str
    metadata: DocumentMetadata
    sections: List[Dict[str, Any]]
    extracted_entities: List[str]
    processing_notes: List[str]

class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self, data_dir: str = "data/documents"):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(parents=True, exist_ok=True)
        
        # 支持的文件类型
        self.supported_types = {
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.html': self._process_html,
            '.htm': self._process_html,
            '.csv': self._process_csv,
            '.json': self._process_json,
            '.xml': self._process_xml
        }
    
    def process_document(self, file_path: Union[str, Path]) -> ProcessedDocument:
        """处理单个文档"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 检测文件类型
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_types:
            raise ValueError(f"不支持的文件类型: {file_ext}")
        
        # 获取文件元数据
        metadata = self._extract_metadata(file_path)
        
        # 处理文档内容
        processor = self.supported_types[file_ext]
        content, sections, notes = processor(file_path)
        
        # 提取实体（简单版本）
        entities = self._extract_basic_entities(content)
        
        return ProcessedDocument(
            content=content,
            metadata=metadata,
            sections=sections,
            extracted_entities=entities,
            processing_notes=notes
        )
    
    def process_directory(self, directory_path: Union[str, Path]) -> List[ProcessedDocument]:
        """批量处理目录中的文档"""
        directory_path = Path(directory_path)
        processed_docs = []
        
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_types:
                try:
                    doc = self.process_document(file_path)
                    processed_docs.append(doc)
                    print(f"✅ 处理完成: {file_path.name}")
                except Exception as e:
                    print(f"❌ 处理失败: {file_path.name} - {e}")
        
        return processed_docs
    
    def _extract_metadata(self, file_path: Path) -> DocumentMetadata:
        """提取文件元数据"""
        stat = file_path.stat()
        
        return DocumentMetadata(
            filename=file_path.name,
            file_type=file_path.suffix.lower(),
            file_size=stat.st_size,
            created_at=str(stat.st_ctime),
            modified_at=str(stat.st_mtime)
        )
    
    def _process_text(self, file_path: Path) -> tuple:
        """处理纯文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk') as f:
                content = f.read()
        
        # 简单分段
        sections = self._split_into_sections(content)
        
        return content, sections, ["文本文件处理完成"]
    
    def _process_markdown(self, file_path: Path) -> tuple:
        """处理Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 转换为HTML然后提取纯文本
        try:
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            content = soup.get_text()
        except:
            content = md_content  # 降级处理
        
        # 提取Markdown结构
        sections = self._extract_markdown_sections(md_content)
        
        return content, sections, ["Markdown文件处理完成"]
    
    def _process_pdf(self, file_path: Path) -> tuple:
        """处理PDF文件"""
        content = ""
        sections = []
        notes = []
        
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    content += page_text + "\n"
                    
                    sections.append({
                        'type': 'page',
                        'title': f'第{i+1}页',
                        'content': page_text,
                        'page_number': i + 1
                    })
                
                notes.append(f"PDF处理完成，共{len(reader.pages)}页")
                
        except Exception as e:
            notes.append(f"PDF处理出错: {e}")
            content = f"PDF文件处理失败: {file_path.name}"
        
        return content, sections, notes
    
    def _process_docx(self, file_path: Path) -> tuple:
        """处理Word文档"""
        content = ""
        sections = []
        notes = []
        
        try:
            doc = docx.Document(file_path)
            
            current_section = {"type": "paragraph", "content": "", "title": "正文"}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content += text + "\n"
                    
                    # 检查是否是标题
                    if para.style.name.startswith('Heading'):
                        if current_section["content"]:
                            sections.append(current_section)
                        current_section = {
                            "type": "heading",
                            "title": text,
                            "content": "",
                            "level": int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                        }
                    else:
                        current_section["content"] += text + "\n"
            
            if current_section["content"]:
                sections.append(current_section)
            
            notes.append("Word文档处理完成")
            
        except Exception as e:
            notes.append(f"Word文档处理出错: {e}")
            content = f"Word文档处理失败: {file_path.name}"
        
        return content, sections, notes
    
    def _process_doc(self, file_path: Path) -> tuple:
        """处理旧版Word文档"""
        # 简化处理，建议转换为docx格式
        notes = ["旧版Word文档(.doc)建议转换为.docx格式以获得更好的处理效果"]
        content = f"旧版Word文档: {file_path.name}\n请转换为.docx格式后重新处理"
        sections = []
        
        return content, sections, notes
    
    def _process_html(self, file_path: Path) -> tuple:
        """处理HTML文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 提取纯文本
        content = soup.get_text()
        
        # 提取结构化内容
        sections = []
        
        # 提取标题和段落
        for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div']):
            if tag.name.startswith('h'):
                sections.append({
                    'type': 'heading',
                    'title': tag.get_text().strip(),
                    'level': int(tag.name[1]),
                    'content': ''
                })
            elif tag.get_text().strip():
                sections.append({
                    'type': 'paragraph',
                    'content': tag.get_text().strip(),
                    'title': ''
                })
        
        return content, sections, ["HTML文件处理完成"]
    
    def _process_csv(self, file_path: Path) -> tuple:
        """处理CSV文件"""
        try:
            df = pd.read_csv(file_path)
            
            # 转换为文本描述
            content = f"CSV文件: {file_path.name}\n"
            content += f"行数: {len(df)}, 列数: {len(df.columns)}\n"
            content += f"列名: {', '.join(df.columns)}\n\n"
            
            # 添加数据样本
            content += "数据样本:\n"
            content += df.head().to_string()
            
            # 创建sections
            sections = [{
                'type': 'data_table',
                'title': 'CSV数据',
                'content': df.to_dict('records'),
                'columns': list(df.columns),
                'row_count': len(df)
            }]
            
            notes = [f"CSV文件处理完成，包含{len(df)}行数据"]
            
        except Exception as e:
            content = f"CSV文件处理失败: {e}"
            sections = []
            notes = [f"CSV处理出错: {e}"]
        
        return content, sections, notes
    
    def _process_json(self, file_path: Path) -> tuple:
        """处理JSON文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 转换为文本描述
        content = f"JSON文件: {file_path.name}\n"
        content += json.dumps(data, ensure_ascii=False, indent=2)
        
        sections = [{
            'type': 'json_data',
            'title': 'JSON数据',
            'content': data
        }]
        
        return content, sections, ["JSON文件处理完成"]
    
    def _process_xml(self, file_path: Path) -> tuple:
        """处理XML文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            xml_content = f.read()
        
        soup = BeautifulSoup(xml_content, 'xml')
        content = soup.get_text()
        
        sections = [{
            'type': 'xml_data',
            'title': 'XML数据',
            'content': xml_content
        }]
        
        return content, sections, ["XML文件处理完成"]
    
    def _split_into_sections(self, content: str) -> List[Dict]:
        """将文本分割为段落"""
        paragraphs = content.split('\n\n')
        sections = []
        
        for i, para in enumerate(paragraphs):
            para = para.strip()
            if para:
                sections.append({
                    'type': 'paragraph',
                    'title': f'段落{i+1}',
                    'content': para
                })
        
        return sections
    
    def _extract_markdown_sections(self, md_content: str) -> List[Dict]:
        """提取Markdown结构"""
        sections = []
        lines = md_content.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if line.startswith('#'):
                # 保存前一个section
                if current_section:
                    sections.append(current_section)
                
                # 创建新section
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()
                current_section = {
                    'type': 'heading',
                    'title': title,
                    'level': level,
                    'content': ''
                }
            elif current_section and line:
                current_section['content'] += line + '\n'
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _extract_basic_entities(self, content: str) -> List[str]:
        """提取基本实体（简单版本）"""
        entities = []
        
        # 提取大写开头的词组（可能是专有名词）
        pattern = r'\b[A-Z][a-zA-Z_]*(?:\s+[A-Z][a-zA-Z_]*)*\b'
        matches = re.findall(pattern, content)
        
        # 过滤常见词汇
        stop_words = {'The', 'This', 'That', 'These', 'Those', 'And', 'Or', 'But', 'In', 'On', 'At', 'To', 'For'}
        entities = [match for match in matches if match not in stop_words and len(match) > 2]
        
        # 去重并限制数量
        entities = list(set(entities))[:50]
        
        return entities
    
    def get_processing_statistics(self, processed_docs: List[ProcessedDocument]) -> Dict:
        """获取处理统计信息"""
        if not processed_docs:
            return {}
        
        total_docs = len(processed_docs)
        total_size = sum(doc.metadata.file_size for doc in processed_docs)
        total_words = sum(doc.metadata.word_count or 0 for doc in processed_docs)
        
        file_types = {}
        for doc in processed_docs:
            file_type = doc.metadata.file_type
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            'total_documents': total_docs,
            'total_size_bytes': total_size,
            'total_words': total_words,
            'file_types': file_types,
            'avg_size_per_doc': total_size / total_docs if total_docs > 0 else 0,
            'avg_words_per_doc': total_words / total_docs if total_docs > 0 else 0
        }
    
    def save_processed_documents(self, processed_docs: List[ProcessedDocument], 
                               output_dir: str = "data/processed") -> str:
        """保存处理后的文档"""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        # 保存每个文档
        for doc in processed_docs:
            doc_data = {
                'content': doc.content,
                'metadata': {
                    'filename': doc.metadata.filename,
                    'file_type': doc.metadata.file_type,
                    'file_size': doc.metadata.file_size,
                    'word_count': len(doc.content.split()),
                    'processing_notes': doc.processing_notes
                },
                'sections': doc.sections,
                'extracted_entities': doc.extracted_entities
            }
            
            # 生成输出文件名
            base_name = Path(doc.metadata.filename).stem
            output_file = output_path / f"{base_name}_processed.json"
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(doc_data, f, ensure_ascii=False, indent=2)
        
        # 保存汇总信息
        summary = {
            'processing_time': str(pd.Timestamp.now()),
            'total_documents': len(processed_docs),
            'statistics': self.get_processing_statistics(processed_docs),
            'file_list': [doc.metadata.filename for doc in processed_docs]
        }
        
        summary_file = output_path / "processing_summary.json"
        with open(summary_file, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
        
        return str(output_path)